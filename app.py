import os
import pandas as pd
from datetime import datetime
import io
import streamlit as st
from langchain import PromptTemplate, LLMChain
#from langchain.llms import GooglePalm
from langchain_google_genai import ChatGoogleGenerativeAI

from docx import Document

if 'key' not in st.session_state:
    st.session_state.key = 0

# Configuration
api_key = st.secrets['google_api_key']
st.set_page_config(page_title="TM Report Automation", layout="wide")
def load_data(path: str) -> pd.DataFrame:
    df = pd.read_csv(path, parse_dates=['transaction_date'])
    # convert boolean
    df['prior_ed_completed'] = df['prior_ed_completed'].astype(bool)
    return df

@st.cache_data
def get_data():
    return load_data('sample_data.csv')

# Initialize LLM
MODEL_NAME = "gemini-2.5-flash"
llm = ChatGoogleGenerativeAI(model=MODEL_NAME, google_api_key=api_key, convert_system_message_to_human=True)#GooglePalm()


@st.cache_data
def generate_tm_report(alert: dict, extra_sections=None) -> dict:
    template = """
Alert ID: {alert_id}
Date: {transaction_date}
Amount: {amount} {currency}
Sender: {sender_account}
Receiver: {receiver_account}
Rule Triggered: {rule_triggered}
Description: {description}
Prior Alerts: {prior_alerts_count}
Prior ED Completed: {prior_ed_completed}
Client Risk Rating: {client_risk_rating}
Jurisdiction Risk: {jurisdiction_risk}
Client ID: {client_id}

Generate a Transaction Monitoring report draft with the following sections:
1. Summary of the transaction
2. Risk factors and anomalies
3. Suggested next steps
"""
    # Add extra sections to the prompt if provided
    if extra_sections:
        for idx, section in enumerate(extra_sections, start=4):
            template += f"\n{idx}. {section}"
    prompt = PromptTemplate(
        input_variables=list(alert.keys()),
        template=template
    )
    chain = LLMChain(llm=llm, prompt=prompt)
    output = chain.run(alert)
    # Parse sections
    sections = {
        'summary': output.split('1.')[1].split('2.')[0].strip(),
        'risk_factors': output.split('2.')[1].split('3.')[0].strip(),
        'next_steps': output.split('3.')[1].strip()
    }
    # Parse extra sections if present
    if extra_sections:
        for idx, section in enumerate(extra_sections, start=4):
            key = f'extra_{idx}'
            # Try to extract the section text
            try:
                if idx < 3 + len(extra_sections):
                    sections[key] = output.split(f'{idx}.')[1].split(f'{idx+1}.')[0].strip()
                else:
                    sections[key] = output.split(f'{idx}.')[1].strip()
            except Exception:
                sections[key] = ''
    return sections

# Streamlit UI
st.title('TM Report Automation')

data = get_data()
alert_ids = data['alert_id'].tolist()

selected_id = st.sidebar.selectbox('Select Alert ID', alert_ids)
alert = data[data['alert_id'] == selected_id].iloc[0].to_dict()

st.sidebar.markdown('### Alert Metadata')
for k, v in alert.items():
    st.sidebar.write(f"**{k}**: {v}")

# Sidebar input for extra sections
extra_sections_input = st.sidebar.text_area(
    'Additional Section Titles (one per line, optional)',
    placeholder='e.g.\nRegulatory Considerations\nSupporting Evidence'
)
extra_sections = [s.strip() for s in extra_sections_input.split('\n') if s.strip()]

if st.sidebar.button('Generate Draft', type="primary") or st.session_state.key == 0:
    report = generate_tm_report(alert, extra_sections=extra_sections)
    # Editable sections
    summary = report['summary']
    risk = report['risk_factors']
    steps = report['next_steps']
    #summary = st.text_area('1. Summary of the transaction', report['summary'], height=150)
    #risk = st.text_area('2. Risk factors and anomalies', report['risk_factors'], height=150)
    #steps = st.text_area('3. Suggested next steps', report['next_steps'], height=150)

    # Extra editable sections
    extra_contents = []
    for idx, section in enumerate(extra_sections, start=4):
        key = f'extra_{idx}'
        #content = st.text_area(f'{idx}. {section}', report.get(key, ''), height=150)
        content = report.get(key, '')
        extra_contents.append((section, content))

    # Download Word doc
    doc = Document()
    doc.add_heading(f'TM Report for Alert {selected_id}', level=1)
    doc.add_heading('1. Summary of the transaction', level=2)
    doc.add_paragraph(summary)
    doc.add_heading('2. Risk factors and anomalies', level=2)
    doc.add_paragraph(risk)
    doc.add_heading('3. Suggested next steps', level=2)
    doc.add_paragraph(steps)
    for idx, (section, content) in enumerate(extra_contents, start=4):
        doc.add_heading(f'{idx}. {section}', level=2)
        doc.add_paragraph(content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    

    # Show non-editable report preview
    st.markdown("---")
    st.subheader("Report Preview (Read-only)")
    st.markdown(f"### TM Report for Alert {selected_id}")
    st.markdown("#### 1. Summary of the transaction")
    st.write(summary)
    st.markdown("#### 2. Risk factors and anomalies")
    st.write(risk)
    st.markdown("#### 3. Suggested next steps")
    st.write(steps)
    for idx, (section, content) in enumerate(extra_contents, start=4):
        st.markdown(f"#### {idx}. {section}")
        st.write(content)

    st.download_button(
        label='Download Report as Word',
        data=buffer,
        file_name=f'TM_Report_{selected_id}.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        type="primary"
    )
    st.session_state.key = 1